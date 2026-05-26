#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Generate Word report for Experiment 13: Logistic Regression

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

PROJECT_ROOT = Path('/home/yfn/repos/machine-learning-labs/lab13_LogisticRegression')
FIG_DIR = PROJECT_ROOT / 'figures' / 'figures'
TEMPLATE_PATH = Path('/mnt/e/01大二下文件夹/机器学习实验报告/第十三次实验/机器学习实验13：逻辑回归算法原生实现的实验报告模板.docx')
OUTPUT_DIR = Path('/mnt/e/01大二下文件夹/机器学习实验报告/第十三次实验')
OUTPUT_PATH = OUTPUT_DIR / '机器学习实验13：数字运营用户流失预测场景下的逻辑回归算法原生实现实验报告.docx'

NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def set_cn_font(run, font_name='宋体'):
    rPr = run._element.find(f'{{{NS}}}rPr')
    if rPr is None:
        rPr = etree.SubElement(run._element, f'{{{NS}}}rPr')
    rFonts = rPr.find(f'{{{NS}}}rFonts')
    if rFonts is None:
        rFonts = etree.SubElement(rPr, f'{{{NS}}}rFonts')
    rFonts.set(f'{{{NS}}}eastAsia', font_name)


def add_normal(doc, text, indent=True, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    if bold:
        run.bold = True
    set_cn_font(run)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    set_cn_font(run)
    return p


def add_image(doc, img_path, caption, width=5.5):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_img.add_run()
    r.add_picture(str(img_path), width=Inches(width))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.first_line_indent = Cm(0)
    run = p_cap.add_run(caption)
    run.font.size = Pt(10.5)
    set_cn_font(run)
    return p_img, p_cap


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Normal Table'
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else etree.SubElement(tbl, f'{{{NS}}}tblPr')
    borders = etree.SubElement(tblPr, f'{{{NS}}}tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = etree.SubElement(borders, f'{{{NS}}}{edge}')
        el.set(f'{{{NS}}}val', 'single')
        el.set(f'{{{NS}}}sz', '4')
        el.set(f'{{{NS}}}space', '0')
        el.set(f'{{{NS}}}color', '000000')
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10.5)
        set_cn_font(run)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(10.5)
            set_cn_font(run)
    return table


def add_tcap(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.bold = True
    set_cn_font(run)
    return p


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Page setup from template
    tmpl = Document(str(TEMPLATE_PATH))
    ts = tmpl.sections[0]
    ds = doc.sections[0]
    ds.page_width = ts.page_width
    ds.page_height = ts.page_height
    for attr in ['left_margin', 'right_margin', 'top_margin', 'bottom_margin']:
        setattr(ds, attr, getattr(ts, attr))
    ds.header.is_linked_to_previous = False
    ds.footer.is_linked_to_previous = False

    #########################
    # TITLE
    #########################
    t = doc.add_heading(
        '机器学习实验13：数字运营用户流失预测场景下的逻辑回归算法原生实现', level=1
    )
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #########################
    # 实验概述
    #########################
    doc.add_heading('实验概述', level=2)

    add_normal(doc,
        '本实验旨在学习和掌握逻辑回归（Logistic Regression）算法的基本原理与原生实现方法。'
        '逻辑回归是一种被广泛用于二分类问题的广义线性模型，以Sigmoid函数为核心将线性回归输出映射到[0,1]区间从而输出概率，'
        '以交叉熵损失函数衡量预测偏差，并以梯度下降法优化模型参数。本实验以数字运营用户流失预测为实际场景，'
        '基于用户行为模拟数据，完全原生手写实现逻辑回归算法，不调用sklearn封装的LogisticRegression模块，'
        '从数据预处理、梯度求解到模型更新的全流程自研实现。')

    add_normal(doc,
        '用户流失预测是数字运营领域的核心命题\u2014\u2014获取一个新用户的成本是留住一个老用户的5-25倍，'
        '因此及时识别高风险流失用户并制定挽留策略具有极高的ROI。逻辑回归在此场景中具有天然优势：特征权重可直接解读，'
        '让运营人员清晰理解\u201c哪些因素最能影响用户流失\u201d；输出连续概率值便于按风险等级分层处理；'
        '计算效率高，可应对大规模用户数据。通过本实验，将建立\u201c可解释性优先、数据驱动决策\u201d的模型价值观。')

    #########################
    # 二、实验目标
    #########################
    doc.add_heading('二、实验目标', level=2)

    targets = [
        '理解逻辑回归的核心原理：Sigmoid函数、对数损失（交叉熵损失）、梯度下降、特征归一化与正则化。',
        '完全原生手写实现逻辑回归算法，不调用sklearn封装的LogisticRegression模块，掌握从数据预处理、梯度求解到模型更新的全流程。',
        '基于数字运营用户行为数据集，完成用户流失预测二分类任务，输出评估指标并分析模型性能。',
        '系统对比逻辑回归与sklearn标准实现的差异，验证原生实现的正确性。',
        '基于逻辑回归输出的特征权重，识别影响用户流失的关键因素并进行业务解读，为用户留存策略提供可落地的数据支撑。',
        '掌握分类阈值调整方法，理解精确率与召回率之间的权衡关系，针对业务需求优化模型决策边界。',
    ]
    for t in targets:
        add_normal(doc, t)

    #########################
    # 三、实验项目
    #########################
    doc.add_heading('三、实验项目', level=2)

    projects = [
        '使用Python与NumPy原生实现逻辑回归算法，包括Sigmoid函数、交叉熵损失函数、小批量梯度下降（MBGD）优化器、L1/L2正则化等核心组件。',
        '构造并预处理数字运营用户行为数据集，包含注册时长、活跃度、消费、互动、满意度五个维度的16个基础特征，并构造3个业务衍生特征。',
        '完成模型训练、预测与评估，输出准确率、精确率、召回率、F1、AUC等指标，绘制损失下降曲线、ROC曲线和混淆矩阵。',
        '与sklearn的逻辑回归标准实现进行对比验证，确保原生算法的正确性。',
        '进行特征权重分析和流失风险分层，输出风险等级分布和对应挽留策略建议，完成业务价值解读和ROI估算。',
    ]
    for t in projects:
        add_normal(doc, t)

    #########################
    # 四、实验内容及实现
    #########################
    doc.add_heading('四、实验内容及实现', level=2)

    # ---- 4.1 ----
    add_h3(doc, '4.1 数据集生成')

    add_normal(doc,
        '实验使用数字运营用户行为模拟数据集，共包含800个样本，每个样本具有16个基础特征，涵盖五个维度：'
        '用户基本特征（注册时长）、活跃特征（最近活跃天数、活跃频率、访问时长、使用功能数）、'
        '消费特征（累计消费金额、最近消费金额、消费频率、客单价）、'
        '互动特征（评论数、分享数、点赞数、反馈次数、客服咨询次数）和'
        '满意度特征（投诉次数、评分均值）。'
        '预测目标为\u201c是否流失\u201d，其中0表示留存、1表示流失。')

    add_normal(doc,
        '数据集生成时按照业务逻辑构建了流失风险分数：最近活跃天数越大、投诉越多、评分越低、活跃越低则流失风险越高。'
        '通过设置分位数阈值控制流失比例约为25%，符合指导书建议的10%-30%范围。'
        '数据集中还按照约1.8%的比例注入了随机缺失值，用于测试预处理流程的鲁棒性。')

    add_table(doc,
        ['维度', '特征名称', '数量'],
        [
            ['基本特征', '注册时长', '1'],
            ['活跃特征', '最近活跃天数、活跃频率、访问时长、使用功能数', '4'],
            ['消费特征', '累计消费金额、最近消费金额、消费频率、客单价', '4'],
            ['互动特征', '评论数、分享数、点赞数、反馈次数、客服咨询次数', '5'],
            ['满意度特征', '投诉次数、评分均值', '2'],
            ['衍生特征', '消费活跃比、互动总量、投诉咨询比', '3'],
        ])
    add_tcap(doc, '表1  数据集特征维度说明')

    add_image(doc, FIG_DIR / 'churn_distribution.png', '图1  用户流失分布', width=4.5)

    # ---- 4.2 ----
    add_h3(doc, '4.2 数据预处理')

    add_normal(doc, '数据预处理是影响逻辑回归训练效果和特征权重可信度的关键步骤，具体流程如下：')
    add_normal(doc, '（1）缺失值处理：对连续特征使用中位数填充，中位数相比均值对异常值更稳健。')
    add_normal(doc,
        '（2）异常值处理：采用IQR（四分位距）法识别异常值，'
        '对超出Q1-1.5\u00d7IQR或Q3+1.5\u00d7IQR的数值进行截断处理（Winsorizing），避免异常值过度影响梯度方向。')
    add_normal(doc,
        '（3）特征工程：对消费金额类（累计消费金额、最近消费金额、客单价）和计数类（评论数、分享数、点赞数）特征'
        '做log1p变换以缓解右偏分布；构造\u201c消费活跃比\u201d（消费频率/活跃频率）、'
        '\u201c互动总量\u201d（评论数+分享数+点赞数）、\u201c投诉咨询比\u201d（投诉次数/客服咨询次数）'
        '三个业务衍生特征，最终特征维度扩展至19维。')
    add_normal(doc,
        '（4）数据划分：按7:3比例分层抽样划分训练集（560样本）和测试集（240样本），确保训练和测试集中流失比例保持一致。')
    add_normal(doc,
        '（5）特征归一化：对所有特征进行Z-score标准化（减去均值、除以标准差），'
        '且仅在训练集上学习均值和标准差后应用到测试集，严格避免数据泄露。'
        '逻辑回归对特征尺度非常敏感，归一化后各特征权重才能直接比较大小。')

    # ---- 4.3 ----
    add_h3(doc, '4.3 逻辑回归算法原理与原生实现')

    add_normal(doc,
        '逻辑回归的核心思想是用线性回归的输出通过Sigmoid函数映射到[0,1]区间，作为概率估计。'
        'Sigmoid函数形式为\u03c3(z) = 1/(1+e^{-z})，其中z = w\u00b7x + b是线性输出，w为权重向量，b为偏置。'
        '当\u03c3(z)接近1时预测为流失，接近0时预测为留存。')

    add_normal(doc,
        '损失函数采用交叉熵损失（对数损失）：L = -[y\u00b7log(\u0177) + (1-y)\u00b7log(1-\u0177)]，'
        '衡量预测概率与真实标签的差距。优化算法采用小批量梯度下降（MBGD），'
        '每次随机抽取batch_size=32个样本计算梯度并更新参数，在计算效率和收敛稳定性之间取得平衡。')

    add_normal(doc,
        '为防止过拟合，引入L2正则化项（Ridge），在损失函数中加入\u03bb\u00b7||w||\u00b2/2，'
        '同时梯度更新时加入\u03bb\u00b7w项。正则化强度\u03bb=0.01，用于控制模型复杂度，防止过度解读个别样本噪声。')

    add_normal(doc,
        '为处理类别不平衡问题（流失用户约占25%），实现类别平衡权重（class_weight=\'balanced\'），'
        '正类（流失）权重为n/(2\u00b7n_pos)，负类（留存）权重为n/(2\u00b7n_neg)，使模型更关注少数类样本。')

    add_normal(doc,
        '原生实现的核心类NativeLogisticRegression封装了以下方法：_sigmoid（Sigmoid函数）、'
        '_compute_loss（交叉熵损失计算）、fit（训练，含小批量梯度下降循环）、predict_proba（输出概率）、'
        'predict（按阈值分类）和get_feature_weights（特征权重分析）。代码完全基于numpy实现，'
        '不依赖sklearn的LogisticRegression封装。')

    # ---- 4.4 ----
    add_h3(doc, '4.4 模型训练与评估')

    add_normal(doc,
        '使用560个训练样本对NativeLogisticRegression模型进行训练，超参数设置为：学习率0.05、'
        '迭代次数1000、batch_size=32、L2正则化、\u03bb=0.01、类别平衡权重。'
        '训练耗时约0.54秒，每100轮迭代输出一次损失值，最终损失收敛至0.2459。')

    add_normal(doc, '在240个测试样本上的评估结果如下表所示：')

    add_table(doc,
        ['指标', '数值', '说明'],
        [
            ['训练集准确率', '0.9125', '模型在训练集上的整体预测准确率'],
            ['测试集准确率', '0.8833', '模型在测试集上的整体预测准确率'],
            ['精确率（Precision）', '0.7051', '预测为流失的用户中真正流失的比例'],
            ['召回率（Recall）', '0.9167', '真实流失用户中被正确识别的比例'],
            ['F1分数', '0.7971', '精确率和召回率的调和平均值'],
            ['AUC分数', '0.9584', 'ROC曲线下面积，反映模型整体判别能力'],
        ])
    add_tcap(doc, '表2  原生逻辑回归模型评估结果')

    add_normal(doc,
        '混淆矩阵显示：真正类（留存\u2192留存）157个，假正类（留存\u2192流失）23个；'
        '假负类（流失\u2192留存）5个，真负类（流失\u2192流失）55个。'
        '模型对流失用户的召回率达到91.67%，说明能有效识别绝大多数高风险用户；'
        '精确率为70.51%，意味着预测为流失的用户中约70%确实流失，存在一定的误报。')

    add_image(doc, FIG_DIR / 'confusion_matrix.png', '图3  混淆矩阵')

    add_image(doc, FIG_DIR / 'loss_history.png', '图2  训练损失下降曲线')

    add_normal(doc,
        '损失从初始值平稳下降并收敛至0.2459，说明学习率和迭代次数的设置合理，模型成功收敛。'
        'ROC曲线（图4）下的面积AUC达到0.9584，表明模型具有良好的整体判别能力，能够有效区分留存用户和流失用户。')

    add_image(doc, FIG_DIR / 'roc_curve.png', '图4  ROC曲线')

    # ---- 4.5 ----
    add_h3(doc, '4.5 分类阈值对比分析')

    add_normal(doc, '默认分类阈值为0.5，在实际业务中可根据需求调整以平衡精确率和召回率。'
               '实验对比了0.30至0.70五个阈值的性能表现：')

    add_table(doc,
        ['阈值', '准确率', '精确率', '召回率', 'F1', '预测流失人数'],
        [
            ['0.30', '0.8375', '0.6129', '0.9500', '0.7451', '93'],
            ['0.40', '0.8625', '0.6588', '0.9333', '0.7724', '85'],
            ['0.50', '0.8833', '0.7051', '0.9167', '0.7971', '78'],
            ['0.60', '0.9000', '0.7727', '0.8500', '0.8095', '66'],
            ['0.70', '0.9208', '0.8596', '0.8167', '0.8376', '57'],
        ])
    add_tcap(doc, '表3  不同分类阈值下的指标对比')

    add_normal(doc,
        '分析发现：随着阈值提高，精确率从61.29%上升至85.96%，但召回率从95.00%下降至81.67%。'
        '若运营目标是\u201c宁可错杀一千，不放过一个\u201d（如高价值用户流失），应选择较低阈值；'
        '若更关注\u201c不打扰留存用户\u201d（精确率优先），应选择较高阈值。'
        'F1在阈值为0.70时达到最高0.8376，表明在该阈值下精确率和召回率取得最佳平衡。')

    add_image(doc, FIG_DIR / 'threshold_compare.png', '图5  不同分类阈值下的精确率、召回率与F1对比')

    # ---- 4.6 ----
    add_h3(doc, '4.6 特征权重分析与业务解读')

    add_normal(doc,
        '逻辑回归最大的优势在于可解释性。特征权重w可直接解读：权重>0表示该特征值越大流失概率越高（正相关），'
        '权重<0表示该特征值越大流失概率越低（负相关），绝对值越大表示对流失的影响越强。'
        '归一化后各特征权重可直接比较。')

    add_normal(doc, '权重绝对值排名前五的特征如下表所示：')

    add_table(doc,
        ['特征', '权重', '影响方向', '业务含义'],
        [
            ['最近活跃天数', '1.4013', '正相关（增加流失风险）', '最近越不活跃，流失风险越高，是核心预警指标'],
            ['客服咨询次数', '1.1423', '正相关（增加流失风险）', '频繁咨询客服的用户往往面临使用困难或不满'],
            ['投诉次数', '1.0666', '正相关（增加流失风险）', '投诉是流失的强信号，投诉越多流失概率越高'],
            ['评分均值', '-0.9734', '负相关（降低流失风险）', '评分越高表示满意度越高，流失风险越低'],
            ['反馈次数', '0.7479', '正相关（增加流失风险）', '频繁反馈问题的用户满意度较低'],
        ])
    add_tcap(doc, '表4  影响用户流失的前5个关键特征')

    add_normal(doc, '业务解读示例：')
    add_normal(doc,
        '\u2022 \u201c最近活跃天数\u201d权重为1.4013，是影响流失最强的特征。'
        '这意味着用户距离上次活跃的天数每增加1个单位，流失的对数几率增加1.40，流失概率显著上升。'
        '运营策略上应重点关注活跃度下降的用户，在沉默早期进行唤醒干预。')
    add_normal(doc,
        '\u2022 \u201c投诉次数\u201d权重为1.0666，说明投诉是流失的强信号。'
        '每增加1次投诉，流失概率大幅提升。运营团队应建立投诉快速响应机制，在用户投诉后及时跟进解决。')
    add_normal(doc,
        '\u2022 \u201c评分均值\u201d权重为-0.9734，说明满意度是留存的关键保护因素。'
        '评分越高，用户流失风险越低，应持续关注低评分用户的反馈并改善产品体验。')

    add_image(doc, FIG_DIR / 'feature_weights.png', '图6  Top12特征权重可视化')

    # ---- 4.7 ----
    add_h3(doc, '4.7 流失风险分层与运营策略')

    add_normal(doc, '基于模型输出的流失概率，将测试集240个用户分为5个风险等级，并制定对应的挽留策略：')

    add_table(doc,
        ['风险等级', '流失概率区间', '人数', '占比', '建议策略'],
        [
            ['极低风险', '0-20%', '137', '57.1%', '持续优化体验，避免过度打扰'],
            ['低风险', '20-40%', '18', '7.5%', '常规关怀，保持内容触达'],
            ['中风险', '40-60%', '19', '7.9%', '发送新功能通知，小额优惠券或任务激励'],
            ['高风险', '60-80%', '19', '7.9%', '定向推送个性化内容，中额优惠券'],
            ['极高风险', '80-100%', '47', '19.6%', '立即专属客服联系，大额优惠券或专属权益'],
        ])
    add_tcap(doc, '表5  用户流失风险等级分布与运营策略')

    add_normal(doc,
        'ROI估算：高风险与极高风险用户共66人（占测试集27.5%），假设挽留成功率为20%，'
        '单个留存用户生命周期价值（LTV）为500元，则预计可创造价值约66\u00d720%\u00d7500=6,600元。')

    add_image(doc, FIG_DIR / 'risk_distribution.png', '图7  用户流失风险等级分布')

    # ---- 4.8 ----
    add_h3(doc, '4.8 与sklearn标准实现对比验证')

    add_normal(doc,
        '为验证原生实现的正确性，使用sklearn的LogisticRegression在相同数据上进行比较'
        '（C=100等价于\u03bb=0.01，max_iter=2000，class_weight=\'balanced\'）：')

    add_table(doc,
        ['模型', '准确率', '精确率', '召回率', 'F1', 'AUC'],
        [
            ['原生逻辑回归', '0.8833', '0.7051', '0.9167', '0.7971', '0.9584'],
            ['sklearn逻辑回归', '0.8875', '0.7324', '0.8667', '0.7939', '0.9524'],
        ])
    add_tcap(doc, '表6  原生实现与sklearn标准实现对比')

    add_normal(doc,
        '对比结果显示：原生实现的各项指标与sklearn高度接近，准确率仅差0.42个百分点，AUC差异为0.006，'
        '验证了原生实现的正确性。细微差异可能来源于优化器实现细节（sklearn使用LIBLINEAR/LBFGS优化器，'
        '而本实验使用小批量梯度下降）以及正则化系数处理方式的不同。'
        '总体来说，原生实现达到了与工业级标准库可比的预测性能。')

    add_image(doc, FIG_DIR / 'model_compare.png', '图8  原生逻辑回归与sklearn逻辑回归对比')

    #########################
    # 五、实验总结与思考
    #########################
    doc.add_heading('五、实验总结与思考', level=2)

    # ---- 5.1 ----
    add_h3(doc, '5.1 实验总结')

    add_normal(doc, '本实验成功原生实现了逻辑回归算法，并在数字运营用户流失预测场景下完成了完整的实验流程。实验结果表明：')
    add_normal(doc,
        '（1）原生实现的逻辑回归在测试集上达到88.33%的准确率、91.67%的召回率和0.9584的AUC，'
        '与sklearn标准实现高度接近，验证了算法实现的正确性。')
    add_normal(doc,
        '（2）特征权重分析揭示了影响用户流失的关键因素：最近活跃天数、投诉次数和客服咨询次数是增加流失风险的核心指标，'
        '评分均值和高活跃度是降低流失风险的保护因素。这些发现高度符合业务直觉，验证了逻辑回归的可解释性优势。')
    add_normal(doc,
        '（3）通过流失风险分层，将测试用户分为5个风险等级，高风险及以上的用户占27.5%，'
        '结合挽留成功率和LTV估算了模型的业务价值。')
    add_normal(doc,
        '（4）分类阈值分析显示，根据业务目标调整阈值可有效平衡精确率和召回率，'
        '在召回率优先的场景（如高价值用户挽留）建议使用较低阈值，'
        '在精确率优先的场景（如减少运营成本）建议使用较高阈值。')

    # ---- 5.2 ----
    add_h3(doc, '5.2 思考题')

    qas = [
        (
            '1. 结合实验结果，说明逻辑回归在用户流失预测场景中的核心优势是什么？为什么可解释性在数字运营中如此重要？',
            '核心优势在于可解释性：特征权重可直接解读，本实验揭示\u201c最近活跃天数\u201d（权重1.4013）和'
            '\u201c投诉次数\u201d（权重1.0666）是流失最强预测因子，可指导运营策略。可解释性之所以重要，'
            '是因为运营人员需要理解\u201c为什么\u201d才能制定针对性挽留策略，而非盲目标记推送。'
            '此外，概率输出便于分层决策，计算效率高可大规模扩展。'
        ),
        (
            '2. 逻辑回归的特征权重可以直接比较大小吗？如果不进行特征归一化，特征权重的解读会有什么问题？',
            '只有归一化后才能直接比较。未归一化时，大尺度特征的权重被压缩、小尺度特征权重被放大，'
            '无法反映真实重要性。例如\u201c累计消费金额\u201d（数千量级）和\u201c评分均值\u201d（1-5），'
            '前者权重会异常偏小但并非不重要。本实验用Z-score标准化确保了权重可比性。'
        ),
        (
            '3. 在用户流失预测中，为什么召回率通常比准确率更重要？如何通过调整分类阈值来平衡精确率和召回率？',
            '漏掉流失用户（假负例）意味着完全失去其未来价值，而误报留存用户（假正例）仅增加少量运营成本，'
            '因此召回率更重要。阈值分析显示：0.30时召回率95.00%（精确率61.29%），0.70时精确率85.96%（召回率81.67%）。'
            '高价值用户挽留用低阈值，成本敏感场景用高阈值。'
        ),
        (
            '4. L1正则化和L2正则化有什么区别？在流失预测场景中，什么时候应该用L1，什么时候应该用L2？',
            'L1（Lasso）加权重绝对值之和，可将不重要特征权重压缩为0，具备特征选择能力；'
            'L2（Ridge）加权重平方和，使权重整体变小但不为0。特征多且有冗余时用L1，'
            '特征间存在多重共线性时用L2更稳定。本实验使用L2（\u03bb=0.01），因19个特征均经过业务筛选。'
        ),
        (
            '5. 假设你是运营人员，拿到逻辑回归输出的特征权重，你会如何利用这些信息来优化用户留存策略？请给出3条具体的、可落地的建议。',
            '建议一：建立活跃度下降预警机制。\u201c最近活跃天数\u201d权重最大，设置沉默天数阈值（如7天），'
            '触发时自动发送唤醒推送或优惠券。\n'
            '建议二：建立投诉快速响应闭环。\u201c投诉次数\u201d是强信号，用户投诉后自动标记高风险、'
            '优先分配客服、24小时内跟进并发送满意度回访。\n'
            '建议三：低评分用户专项运营。\u201c评分均值\u201d是保护因素，对评分<3.5的用户'
            '定期推送改进说明和新功能邀请，给予专属优惠券提升满意度。'
        ),
        (
            '6. 逻辑回归的决策边界是线性的，在实际业务中可能不够灵活。有什么方法可以让逻辑回归也能捕捉非线性关系？',
            '（1）特征交叉：如\u201c活跃频率\u00d7消费金额\u201d捕捉交互效应。'
            '（2）多项式特征：平方、立方等变换，如\u201c评分均值\u00b2\u201d。本实验的log1p变换和比例特征已是实践。'
            '（3）分箱：将连续特征离散化为分段哑变量。（4）核技巧：如RBF核将特征映射到高维空间。'
        ),
        (
            '7. 讨论小批量梯度下降、批量梯度下降、随机梯度下降的优缺点。在流失预测场景中，你会选择哪种，为什么？',
            'BGD用全部样本，方向准但计算大；SGD用单样本，快但不稳；MBGD用一批样本，平衡速度和准确性。'
            '推荐MBGD：用户数据常达百万级，BGD无法实时更新，SGD不稳定难调参。'
            '本实验用batch_size=32的MBGD，0.54秒完成1000轮迭代，效率良好。'
        ),
        (
            '8. 用户流失预测通常面临类别不平衡问题（流失用户远少于留存用户）。这会对逻辑回归的训练产生什么影响？有什么解决方法？',
            '不平衡导致模型偏向多数类，召回率极低。解决方法：（1）类别权重，本实验用class_weight=\'balanced\'。'
            '（2）SMOTE过采样或随机欠采样。（3）阈值调整，如本实验阈值从0.5降至0.3时召回率从91.67%升至95.00%。'
            '（4）用AUC代替准确率评估。'
        ),
        (
            '9. 有人说\u201c逻辑回归太简单了，不如用深度学习\u201d，结合本次实验，谈谈你对这句话的看法？在什么场景下应该用简单模型，什么场景下应该用复杂模型？',
            '逻辑回归简单但不可替代：高可解释性、低计算成本（0.54秒）、数据高效（560样本）。'
            '简单模型适合数据量小、需可解释性、资源有限的场景；'
            '复杂模型适合数据量极大、存在复杂非线性关系、不要求可解释性的场景。'
            '应遵循\u201c奥卡姆剃刀\u201d原则，从简单模型开始。'
        ),
        (
            '10. 假设你要部署这个流失预测模型到线上，每周更新一次，你需要考虑哪些工程和业务问题？',
            '工程方面：（1）自动化ETL数据管道；（2）MLOps自动化训练评估上线；'
            '（3）监控数据漂移和指标衰减；（4）高可用API推理服务；（5）真实流失标签反馈闭环。\n'
            '业务方面：（1）运营团队培训；（2）A/B测试验证效果；（3）数据合规与审计日志；（4）定期ROI核算。'
        ),
    ]

    for q, a in qas:
        add_normal(doc, q, bold=True)
        add_normal(doc, a)

    # Save
    doc.save(str(OUTPUT_PATH))
    print(f'Document saved to: {OUTPUT_PATH}')


if __name__ == '__main__':
    main()
