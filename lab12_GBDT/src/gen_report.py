# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

FIG_BASE = "/mnt/e/01大二下文件夹/机器学习实验报告/第十二次实验/figures"

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_heading_cn(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h


def add_para(text, bold=False, size=11, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.first_line_indent = Pt(22)
    return p


def add_para_ni(text, bold=False, size=11, align=None, space_after=6):
    """Paragraph without indent"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(size)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_figure(fig_name, caption):
    img_path = os.path.join(FIG_BASE, fig_name)
    if not os.path.exists(img_path):
        return
    doc.add_picture(img_path, width=Inches(5.0))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.font.name = '宋体'
    run.font.size = Pt(9)
    run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.space_after = Pt(10)


def set_cell(cell, text, bold=False, size=9):
    cell.text = text
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(size)
            run.bold = bold
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


# ===== 封面 =====
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("机器学习实验12")
run.font.name = '黑体'
run.font.size = Pt(28)
run.bold = True
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("梯度提升树算法原生实现\n（电商直播流量转化率预测场景）")
run.font.name = '黑体'
run.font.size = Pt(20)
run.bold = True
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

for _ in range(4):
    doc.add_paragraph()

info = [
    ("姓    名", "袁丰年"),
    ("学    号", "2024111010327"),
    ("班    级", "2407班"),
    ("日    期", "2026年5月"),
]
for label, value in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{label}：{value}")
    run.font.name = '宋体'
    run.font.size = Pt(14)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# =====================================================================
# 一、实验概述
# =====================================================================
add_heading_cn("一、实验概述", level=1)

add_heading_cn("1.1 实验背景", level=2)
add_para(
    "随着直播电商行业的快速发展，流量转化率预测已成为核心业务命题。"
    "直播电商与传统电商最大的差异在于时效性和即时互动性：一场直播的转化率受到"
    "实时在线人数、互动指标（点赞/评论/分享）、商品讲解时长、主播话术阶段、"
    "库存紧张程度等多维动态因素的共同影响。这些特征之间存在复杂的非线性交互关系，"
    "且某些特征（如互动热度）对转化率的影响具有累积效应——开场冷淡可能导致全程低迷，"
    "而中段爆发可能带动尾盘热销。"
)
add_para(
    "在前序实验中，我们已经通过随机森林算法探索了电商运营数据的分类预测。"
    "随机森林通过Bagging并行集成机制有效抑制了单棵决策树的过拟合问题，但它在处理"
    "具有明确序列依赖性和递进学习需求的场景时存在局限——它无法从错误中持续改进，"
    "而是依赖多树的简单投票。梯度提升树（Gradient Boosting Decision Tree, GBDT）"
    "正是为解决这类问题而设计的。与随机森林的群体智慧不同，GBDT强调精益求精："
    "每一棵新树都专注于纠正前面所有树的错误，通过持续的自我迭代逼近真实规律。"
    "这种递进式学习机制使GBDT天然适合直播转化率预测——它能够捕捉互动指标的累积效应，"
    "理解不同阶段特征对转化的差异化影响，并从历史直播的失败案例中持续学习。"
)

add_heading_cn("1.2 实验目的", level=2)
purposes = [
    "理解梯度提升树的核心原理：串行集成、负梯度拟合、损失函数优化、残差迭代。",
    "完全原生手写实现GBDT算法，不调用sklearn封装的GradientBoosting模块，掌握从损失函数计算、梯度求解、决策树拟合到残差更新的全流程逻辑。",
    "基于电商直播运营数据集，完成不同转化率等级（低转化/普通转化/高转化/爆款转化）的精准预测。",
    "系统对比随机森林和GBDT在直播转化率预测场景中的差异，理解GBDT在处理序列依赖性和递进学习需求场景中的优势。",
    "基于GBDT输出的特征重要性分析，识别影响直播转化率的关键因素，为直播运营策略优化提供数据支撑。",
]
for i, purpose in enumerate(purposes, 1):
    add_para_ni(f"{i}. {purpose}")

add_heading_cn("1.3 实验环境", level=2)
env_items = [
    ("硬件环境：", "计算机（CPU\u2265i5，内存\u22658G）"),
    ("软件环境：", "Python 3.8+、numpy、pandas、scikit-learn、matplotlib"),
    ("数据集：", "电商直播运营数据集（360条样本，13个特征，4类标签）"),
]
for label, value in env_items:
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(11)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run = p.add_run(value)
    run.font.name = '宋体'
    run.font.size = Pt(11)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_heading_cn("1.4 实验预备知识", level=2)
add_para(
    "基础Python编程（数据读取、数组操作、递归函数实现、类封装）；"
    "机器学习基础概念（特征、标签、训练集、测试集、过拟合、泛化能力）；"
    "基础数学知识（梯度、导数、损失函数、最小二乘法）；"
    "梯度提升机制、负梯度拟合、CART回归树、学习率衰减等新增预备知识。"
)

# =====================================================================
# 二、实验目标
# =====================================================================
doc.add_page_break()
add_heading_cn("二、实验目标", level=1)
objectives = [
    "深刻理解梯度提升树算法的核心思想，明确其与随机森林（Bagging）的本质差异。",
    "熟练掌握GBDT的核心概念（负梯度拟合、串行迭代、学习率控制）与完整执行流程。",
    "能够基于Python原生实现完整的GBDT多分类算法，不依赖第三方聚类库的封装接口。",
    "掌握GBDT关键参数（树数量、学习率、最大深度）的科学选择方法，分析参数变化对模型性能的影响规律。",
    "通过电商直播数据实验，验证GBDT对极端样本（爆款直播）的识别能力，对比与随机森林的差异。",
    "基于特征重要性分析，为直播运营策略优化提供数据支撑，建立数据驱动决策的思维。",
]
for i, obj in enumerate(objectives, 1):
    add_para_ni(f"{i}. {obj}")

# =====================================================================
# 三、实验项目
# =====================================================================
add_heading_cn("三、实验项目", level=1)
add_para(
    "本实验基于电商直播运营数据集，使用Python与NumPy原生实现完整的梯度提升树（GBDT）"
    "多分类算法，完成CART回归树实现、损失函数定义、负梯度计算、串行迭代训练、"
    "特征重要性计算等主要步骤。数据集包含13个直播运营特征，预测目标为四类转化率等级"
    "（低转化/普通转化/高转化/爆款转化）。通过对比随机森林的预测效果，分析GBDT的"
    "递进学习机制在直播转化率预测中的优势，并完成超参数对比实验和特征重要性分析。"
)

# =====================================================================
# 四、实验内容及实现
# =====================================================================
doc.add_page_break()
add_heading_cn("四、实验内容及实现", level=1)

# ─── 4.1 数据集与预处理 ───
add_heading_cn("4.1 数据集与预处理", level=2)
add_para(
    "本实验使用电商直播运营数据集，包含360条直播场次的记录，涵盖13个核心运营特征："
    "场次在线峰值、场均观看人数、平均观看时长、点赞数、评论数、分享数、商品点击数、"
    "讲解商品数量、场均GMV、客单价、退货率、粉丝占比、付费流量占比。"
    "转化率等级按业务阈值定义为四类：低转化（0，转化率<2%）、普通转化（1，2%-5%）、"
    "高转化（2，5%-10%）、爆款转化（3，>10%）。数据覆盖不同品类"
    "（服饰/美妆/食品/数码）、不同时段（早场/午场/晚场/深夜场）和不同主播风格"
    "（测评型/福利型/专业讲解型/娱乐互动型），确保模型能学习到泛化性强的规律。"
)
add_para(
    "数据预处理步骤包括："
    "（1）缺失值处理——连续特征使用中位数填充（对极端值更稳健），比例特征（退货率、粉丝占比、付费流量占比）"
    "使用均值填充，缺失率超过30%的特征直接删除；"
    "（2）异常值截断——使用1%和99%分位数对极端值进行截断处理，保留爆款样本但限制其影响范围，"
    "特别注意不将爆款直播的数据视为异常值删除；"
    "（3）数据划分——采用分层抽样按7:3比例划分为训练集和测试集，设置random_state=42保证实验可复现。"
)

add_para(
    "图1展示了数据集中四类转化率等级的样本分布情况。"
    "从图中可以看出：低转化99条（27.5%）、普通转化124条（34.4%）、"
    "高转化98条（27.2%）、爆款转化39条（10.8%）。"
    "四类样本分布相对均衡，爆款转化虽然占比最低但仍有39条（10.8%），"
    "为模型学习极端样本特征提供了足够的数据支持。"
)
add_figure("class_distribution.png", "图1 转化率等级样本分布")

# ─── 4.2 GBDT算法原理与原生实现 ───
add_heading_cn("4.2 GBDT算法原理与原生实现", level=2)
add_para(
    "GBDT（Gradient Boosting Decision Tree）是一种基于Boosting框架的串行集成学习算法。"
    "与随机森林的并行独立训练不同，GBDT的核心思想是通过多轮迭代，每轮训练一棵CART回归树"
    "来拟合当前模型的负梯度（即残差方向），从而逐步逼近真实标签。"
)
add_para(
    "对于多分类任务，本实验采用One-vs-Rest（OvR）策略，为每个类别训练一个独立的二分类器。"
    "每个二分类器的训练流程如下："
    "（1）初始化——将初始预测值设为训练集中该类别的对数几率；"
    "（2）负梯度计算——计算当前预测概率与真实标签的差值（即负梯度/伪残差）；"
    "（3）回归树拟合——用CART回归树拟合负梯度；"
    "（4）步长更新——以学习率缩放树的预测值后累加到当前预测中；"
    "（5）迭代——重复（2）-（4）直到达到指定迭代轮数。"
    "最终将四个二分类器的输出概率进行归一化，得到多分类概率分布。"
)

# ─── 4.3 CART回归树实现 ───
add_heading_cn("4.3 CART回归树实现", level=2)
add_para(
    "CART（Classification and Regression Tree）回归树作为GBDT的基学习器，采用严格的二叉分裂策略。"
    "分裂标准选择均方误差（MSE）减少量最大化：遍历所有特征和候选阈值，计算分裂前后的MSE增益，"
    "选择增益最大的分裂点。"
)
add_para(
    "关键参数设置及其业务含义：最大深度max_depth=3，限制树的复杂度，避免过拟合到特定直播场次；"
    "叶节点最小样本数min_samples_split=6，确保每个分裂规则至少基于6场直播的经验；"
    "候选阈值通过对特征排序后取分位数得到，最多10个候选值以控制计算复杂度。"
    "特征重要性通过累加各特征在树分裂中带来的总增益并归一化得到。"
)

# ─── 4.4 模型训练与评估 ───
add_heading_cn("4.4 模型训练与评估", level=2)
add_para(
    "使用默认参数（n_estimators=20, max_depth=3, learning_rate=0.20, min_samples_split=6）"
    "训练原生GBDT模型。模型在训练过程中每轮迭代的损失变化如图2所示。"
)
add_para(
    "从图2可以看出，随着迭代轮数的增加，训练损失呈快速下降趋势并趋于收敛。"
    "前5轮损失下降最为显著，说明GBDT在早期迭代中能快速纠正主要偏差；"
    "10轮之后损失曲线趋于平缓，表明模型已基本收敛。"
    "这种快速收敛的特性得益于GBDT的负梯度拟合机制——每轮都针对当前最大的预测偏差进行修正。"
)
add_figure("gbdt_loss_curve.png", "图2 GBDT训练损失变化曲线（20轮迭代）")

add_para(
    "模型在测试集上的分类效果通过混淆矩阵直观展示。图3为原生GBDT的混淆矩阵，"
    "图4为随机森林的混淆矩阵，两者在相同测试集上的表现可供对比。"
)
add_para(
    "从GBDT混淆矩阵（图3）可以看出：低转化和普通转化的分类准确率较高，"
    "高转化类别实现了完全正确的分类（召回率100%），但爆款转化出现了3个漏报"
    "（被误判为普通转化或高转化），说明模型对极端样本的识别仍有提升空间。"
    "随机森林混淆矩阵（图4）中，爆款转化仅漏报1个，整体表现更优。"
)
add_figure("gbdt_confusion_matrix.png", "图3 原生GBDT混淆矩阵")
add_figure("rf_confusion_matrix.png", "图4 随机森林混淆矩阵")

# ─── 4.5 对比实验 ───
doc.add_page_break()
add_heading_cn("4.5 对比实验：GBDT vs 随机森林", level=2)
add_para(
    "为保证对比的公平性，两种模型使用完全相同的训练集和测试集（7:3分层抽样），"
    "并采用相近的核心参数配置。随机森林使用n_estimators=20、max_depth=4、"
    "class_weight=balanced，以平衡类别权重。评估指标包括训练准确率、测试准确率、"
    "加权精确率/召回率/F1、爆款召回率、训练时间和预测时间。"
)

# 指标对比表
table = doc.add_table(rows=3, cols=9)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["模型", "训练准确率", "测试准确率", "加权精确率", "加权召回率", "加权F1", "爆款召回率", "训练时间(s)", "预测时间(s)"]
for i, h in enumerate(headers):
    set_cell(table.rows[0].cells[i], h, bold=True, size=8)

gbdt_data = ["原生GBDT", "0.9802", "0.9352", "0.9371", "0.9352", "0.9338", "0.7500", "1.0477", "0.0030"]
rf_data = ["随机森林", "0.9841", "0.9722", "0.9737", "0.9722", "0.9721", "0.9167", "0.0153", "0.0006"]
for i, data in enumerate([gbdt_data, rf_data], 1):
    for j, val in enumerate(data):
        set_cell(table.rows[i].cells[j], val, size=8)

add_para_ni("表1 GBDT与随机森林模型指标对比（测试集）", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

add_para(
    "从表1的对比数据可以看出以下关键发现。"
    "第一，在整体准确率上，随机森林（97.22%）略高于原生GBDT（93.52%），"
    "这是因为数据集规模较小（360条），随机森林的Bagging机制在小样本上具有更好的泛化稳定性。"
    "第二，在爆款召回率方面，随机森林（91.67%）也优于GBDT（75.00%），"
    "这主要是由于本实验GBDT树数量较少（20棵），串行迭代优势尚未充分体现。"
    "第三，GBDT的训练时间（1.05s）远长于随机森林（0.015s），这是串行训练的固有代价，"
    "但在小规模数据集上仍在可接受范围内。"
    "第四，两者的过拟合程度均较低（训练集与测试集准确率差值小于5%），模型泛化能力良好。"
)

add_para(
    "图5更直观地对比了两者在核心指标上的差异，"
    "图6则展示了训练集与测试集准确率的差值（即过拟合程度）。"
)
add_figure("model_compare.png", "图5 GBDT与随机森林核心指标对比")
add_figure("overfitting_compare.png", "图6 过拟合程度对比（训练准确率 - 测试准确率）")

add_para(
    "在爆款识别这一核心任务上，PR曲线能更全面地评估模型在不同阈值下的表现。"
    "图7展示了两种模型对爆款转化（类别3）的PR曲线对比。"
    "从曲线可以看出，随机森林在不同召回率水平下均保持了较高的精确率，"
    "AP值（0.96）高于GBDT（0.82）。"
    "这表明在本实验的数据规模和参数设置下，随机森林对爆款直播的识别更为稳定可靠。"
    "但值得注意的是，GBDT的PR曲线在低召回率区域（0-0.4）表现与随机森林接近，"
    "说明通过调整决策阈值或优化超参数，GBDT在爆款识别上仍有提升潜力。"
)
add_figure("burst_pr_curve_compare.png", "图7 爆款类别PR曲线对比")

# ─── 4.6 特征重要性分析 ───
add_heading_cn("4.6 特征重要性分析", level=2)
add_para(
    "GBDT的一个重要优势是能够输出特征重要性——即各特征在树分裂过程中带来的总增益的归一化权重。"
    "这一指标可以帮助运营团队理解哪些直播指标对转化率预测最为关键，从而制定更有针对性的优化策略。"
)

feat_table = doc.add_table(rows=6, cols=2)
feat_table.style = 'Light Shading Accent 1'
feat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell(feat_table.rows[0].cells[0], "特征", bold=True, size=10)
set_cell(feat_table.rows[0].cells[1], "重要性", bold=True, size=10)

feats = [
    ("商品点击数", "0.4340"),
    ("场均GMV", "0.2693"),
    ("场均观看人数", "0.1025"),
    ("分享数", "0.0905"),
    ("平均观看时长", "0.0455"),
]
for i, (feat, val) in enumerate(feats, 1):
    set_cell(feat_table.rows[i].cells[0], feat, size=10)
    set_cell(feat_table.rows[i].cells[1], val, size=10)

add_para_ni("表2 GBDT特征重要性Top 5", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

add_para(
    "从表2和图8可以看出，特征重要性呈现出明显的长尾分布。"
    "商品点击数以43.40%的权重高居榜首，表明用户的商品点击行为是转化率的最强预测信号——"
    "点击欲望越强，最终转化的概率越高。场均GMV以26.93%位列第二，"
    "反映了实际成交额对转化率等级的决定性作用。"
    "两者合计占比超过70%，是预测转化率的核心特征。"
)
add_para(
    "场均观看人数（10.25%）和分享数（9.05%）分别代表了直播的触达广度和内容传播力，"
    "高观看量和活跃的分享行为通常是高转化的前兆。平均观看时长（4.55%）虽然排名第五，"
    "但反映了内容吸引力对用户决策的影响——停留时间越长，被说服转化的可能性越大。"
    "这些结果明确了运营优化的重点方向：应聚焦于商品展示优化、GMV提升策略和观众互动引导。"
)
add_figure("gbdt_feature_importance.png", "图8 原生GBDT特征重要性排序")

# ─── 4.7 超参数对比实验 ───
add_heading_cn("4.7 超参数对比实验", level=2)
add_para(
    "超参数的选择直接影响GBDT的预测性能和泛化能力。本实验选取了5组具有代表性的超参数组合，"
    "从树数量（n_estimators）、学习率（learning_rate）和最大深度（max_depth）三个维度"
    "系统探究参数变化对模型性能的影响规律。"
)

hp_table = doc.add_table(rows=6, cols=6)
hp_table.style = 'Light Shading Accent 1'
hp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hp_headers = ["参数组合", "n_estimators", "learning_rate", "max_depth", "测试准确率", "爆款召回率"]
for i, h in enumerate(hp_headers):
    set_cell(hp_table.rows[0].cells[i], h, bold=True, size=8)

hp_data = [
    ("树10_lr0.1_深度2", "10", "0.1", "2", "0.8426", "0.0000"),
    ("树20_lr0.1_深度3", "20", "0.1", "3", "0.9259", "0.5833"),
    ("树20_lr0.2_深度3", "20", "0.2", "3", "0.9352", "0.7500"),
    ("树30_lr0.2_深度3", "30", "0.2", "3", "0.9352", "0.7500"),
    ("树20_lr0.2_深度4", "20", "0.2", "4", "0.9444", "0.8333"),
]
for i, row_data in enumerate(hp_data, 1):
    for j, val in enumerate(row_data):
        set_cell(hp_table.rows[i].cells[j], val, size=8)

add_para_ni("表3 GBDT超参数对比实验结果", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

add_para(
    "从表3和图9可以得出以下结论。"
    "第一，树数量从10增加到20时准确率提升显著（84.26%→92.59%），但从20增加到30时准确率持平，"
    "说明20棵树已基本满足该数据集的训练需求，继续增加带来边际收益递减。"
    "第二，学习率从0.1提升到0.2使爆款召回率从58.33%大幅提升至75.00%，"
    "适当增大步长有助于模型更快地逼近最优解，特别是在树数量有限的情况下。"
    "第三，最大深度从3增加到4使测试准确率提升约1%、爆款召回率提升约8%，"
    "但训练时间增加了约50%（1.08s→1.55s），体现了性能与效率的权衡。"
    "第四，深度不足（深度=2时）导致模型严重欠拟合，爆款召回率为0，"
    "说明至少需要3层深度才能捕捉有效的特征交互关系。"
)
add_figure("gbdt_hyperparam_compare.png", "图9 GBDT超参数对性能的影响")

# =====================================================================
# 五、实验总结与思考
# =====================================================================
doc.add_page_break()
add_heading_cn("五、实验总结与思考", level=1)

add_heading_cn("5.1 实验结果总结", level=2)
add_para(
    "本实验成功实现了梯度提升树算法的原生手写实现，并在电商直播流量转化率预测"
    "数据集上完成了完整的实验验证。主要实验结论如下。"
)
summaries = [
    "原生GBDT在测试集上达到93.52%的准确率，爆款召回率为75.00%；随机森林在相同条件下准确率为97.22%，爆款召回率为91.67%。",
    "特征重要性分析表明，商品点击数（43.40%）和场均GMV（26.93%）是影响转化率最重要的两个特征，合计占比超过70%。",
    "超参数优化实验确定最佳组合为n_estimators=20、learning_rate=0.2、max_depth=4，测试准确率94.44%，爆款召回率83.33%。",
    "GBDT的过拟合程度较低（训练集与测试集准确率差值<5%），泛化能力良好；但串行训练时间（1.05s）长于随机森林（0.015s）。",
    "在数据集规模较小时，随机森林凭借Bagging机制在稳定性和准确率上表现更优；GBDT的优势在更大规模和更强序列依赖性的数据上才能充分体现。",
]
for s in summaries:
    add_para_ni(f"\u2022 {s}")

add_heading_cn("5.2 GBDT与随机森林对比分析", level=2)
add_para(
    "在本实验的数据集上，随机森林的整体表现优于原生GBDT。深入分析原因如下："
    "（1）数据集规模较小（360条），随机森林的Bagging机制通过Bootstrap抽样和投票融合"
    "在小样本场景下具有天然的稳定性优势；"
    "（2）GBDT需要足够的迭代轮数和数据量才能充分发挥递进学习的优势，"
    "20棵树的配置在360条数据上可能尚未达到最佳性能边界；"
    "（3）当前数据集的特征维度（13维）相对有限，特征交互的复杂程度不足以完全体现"
    "GBDT在捕捉非线性关系上的优势。"
)
add_para(
    "尽管如此，GBDT在爆款识别上仍展现出潜力——通过调整超参数（深度增至4、学习率提至0.2）"
    "可将爆款召回率从75%提升至83.33%。在更大规模、具有更强序列依赖性的数据集上"
    "（如包含多场次时序特征的直播数据），GBDT的递进学习优势将更加明显。"
    "这一结果也说明，算法选择需要结合具体场景和数据特点，没有绝对的优劣之分。"
)

add_heading_cn("5.3 思考题", level=2)

questions = [
    ("1. 结合实验结果，说明GBDT相对于随机森林在电商直播转化率预测场景中的核心优势是什么？为什么递进学习机制更适合处理直播间的动态特征？",
     "本实验中GBDT测试准确率93.52%，相比于随机森林的97.22%偏低，但GBDT的核心价值在于串行递进学习——每棵树纠正前序错误，从而捕捉特征间的累积效应。"
     "直播间的动态特征（如开场热度影响中期互动、中期互动影响尾盘转化）具有明确的递进依赖关系，GBDT的串行机制天然适配这种场景。"),

    ("2. 超参n_estimators（树数量）和learning_rate（学习率），在直播转化率预测场景中如何合理调优？如何在模型性能和训练时间之间取得平衡？",
     "先固定学习率0.1找到合适的树数量区间（本实验20棵足够，30棵时准确率持平），再调整学习率（0.1升至0.2后爆款召回率从58%提升至75%）。"
     "模型性能与训练时间的平衡：适度增加深度和树数量可提升精度，但边际收益递减，需根据业务需求取舍。"),

    ("3. max_depth（最大深度）参数如何影响模型对特征交互的捕捉能力？如果设为1（决策树桩），GBDT会退化为什么算法？业务上会导致什么问题？",
     "max_depth越大，树能捕捉的特征交互阶数越高。若设为1（决策树桩），GBDT退化为类似AdaBoost的加法模型，每棵树只基于一个特征分裂，"
     "无法捕捉\u201c高在线人数\u00d7高互动率\u201d这类组合特征的协同影响。本实验深度=2时爆款召回率为0，深度=3时升至75%，充分说明深度不足会严重损失预测精度。"),

    ("4. 为什么电商直播（特别是服饰、美妆、食品类目，高互动、高动态、需要精准预测）更适合GBDT，而不是随机森林？",
     "电商直播具有高互动、高动态特征：前期热度影响后期转化，特征间存在复杂非线性交互，爆款极端样本极具价值。"
     "GBDT的串行学习捕捉递进关系，树结构支持特征交互，残差放大机制敏感识别极端样本；随机森林的并行投票难以处理这些特点。"),

    ("5. GBDT的特征重要性评估如何帮助运营团队优化直播策略？请结合实验结果，提出至少3条具体的运营建议。",
     "特征重要性明确了优化方向：（1）商品点击数（43.40%）——优化商品展示方式，如倒计时标签、弹窗样式以提升点击；（2）场均GMV（26.93%）——分析高转化场次GMV构成，提炼选品策略；（3）分享数（9.05%）和观看时长（4.55%）——设计分享抽奖等互动玩法提升留存。"),

    ("6. 结合GBDT的递进学习机制，讨论如何利用\u201c失败直播\u201d（低转化场次）的数据来优化模型，从而避免重复踩坑。",
     "低转化场次的残差较大，后续树会重点关注这些样本，自动学习失败模式。运营可将低转化场次的特征组合作为反模式输入模型，使模型学会预警\u201c高在线人数\u00d7低互动率\u201d等危险信号，从而避免重复踩坑。"),

    ("7. GBDT的\u201c黑盒\u201d特性可能导致可解释性不足，在实际向运营团队汇报时，如何用业务语言解释模型的预测结果？",
     "类比为\u201c不断积累经验的运营分析师团队\u201d：第一位分析师给出初步判断，后续的分析师依次纠正前人的遗漏，迭代越多判断越准。"
     "同时用特征重要性（如商品点击数最重要）解释模型关注点，而非技术术语。"),

    ("8. GBDT仍有哪些局限？未来如何用更复杂的算法（如XGBoost、LightGBM）在保持递进学习能力的同时提升效率和可解释性？",
     "GBDT局限：串行训练慢、对异常值敏感、高维稀疏数据不佳、调参复杂。"
     "XGBoost引入正则化和二阶梯度信息，LightGBM通过单边梯度采样和互斥特征捆绑大幅提速，两者在保持递进学习能力的同时显著提升了效率和可解释性。"),

    ("9. 有人说\u201c直播靠运气，算法没用\u201d，结合本次实验，谈谈你对这句话的看法？算法和运营经验应该如何结合？",
     "实验表明转化率可通过特征预测，算法并非无用。最佳结合方式：算法客观识别关键因素和模式，运营经验解释业务逻辑并落地优化建议，形成数据与经验相互验证的闭环。"),

    ("10. 如果要预测下一场直播的转化率，GBDT应该如何处理时间序列特征（如\u201c上一场的转化率\u201d、\u201c连续3场的平均转化率\u201d）？",
     "通过特征工程引入时序信息：将上一场转化率、连续3场平均转化率、转化率差分等构造为额外数值特征输入。"
     "也可构造时序交叉特征（如上周同期转化率与时段交互）。GBDT无需特殊时序机制，数值输入即可。"),
]

for q, a in questions:
    add_para(q, bold=True)
    add_para(a)

# ===== 尾部签名 =====
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("袁丰年  2024111010327   2407班")
run.font.name = '宋体'
run.font.size = Pt(12)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

output_path = "/mnt/e/01大二下文件夹/机器学习实验报告/第十二次实验/机器学习实验12：梯度提升树算法原生实现的实验报告.docx"
doc.save(output_path)
print(f"文档已保存至: {output_path}")
