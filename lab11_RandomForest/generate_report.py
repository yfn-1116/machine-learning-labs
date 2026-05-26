# -*- coding: utf-8 -*-
"""生成机器学习实验11：随机森林算法实验报告 Word 文档"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

PROJECT_ROOT = Path(__file__).resolve().parent
FIGURES_DIR = PROJECT_ROOT / "figures"
REPORT_PATH = PROJECT_ROOT / "机器学习实验11_随机森林算法实验报告.docx"

# Color scheme
DARK_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
MED_BLUE = RGBColor(0x2C, 0x5F, 0x8A)
HEADER_BG = "1B3A5C"


def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color,
    })
    shading.append(shading_elm)


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = DARK_BLUE if level == 1 else MED_BLUE
    return heading


def add_body_text(doc, text, bold=False, font_size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "微软雅黑"
    run.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(20)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = "Consolas"
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(14)
    return p


def add_image_centered(doc, img_path, width_inches=5.5, caption=""):
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(width_inches))
        if caption:
            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_p.add_run(caption)
            cap_run.font.size = Pt(9)
            cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            cap_run.italic = True


def build_report():
    doc = Document()

    # ========== 全局样式设置 ==========
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ========== 封面 ==========
    for _ in range(4):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("机器学习实验报告")
    run.font.size = Pt(28)
    run.font.color.rgb = DARK_BLUE
    run.bold = True

    doc.add_paragraph()

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run("实验十一：随机森林算法原生实现\n——电商代运营货品结构盘点分类场景")
    run.font.size = Pt(16)
    run.font.color.rgb = MED_BLUE

    for _ in range(3):
        doc.add_paragraph()

    info_items = [
        ("课程名称", "机器学习"),
        ("实验项目", "随机森林算法原生实现"),
        ("实验日期", datetime.date.today().strftime("%Y年%m月%d日")),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}：{value}")
        run.font.size = Pt(14)
        run.font.color.rgb = DARK_BLUE

    doc.add_page_break()

    # ========== 目录页 ==========
    add_heading_styled(doc, "目  录", level=1)
    toc_items = [
        "一、实验目的",
        "二、实验原理",
        "三、实验环境",
        "四、数据集说明",
        "五、算法实现",
        "六、实验内容与步骤",
        "七、实验结果与分析",
        "八、结论与体会",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.color.rgb = MED_BLUE
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ================================================================
    # 一、实验目的
    # ================================================================
    add_heading_styled(doc, "一、实验目的", level=1)
    objectives = [
        "理解随机森林（Random Forest）算法的基本原理和核心思想，包括 Bagging 集成策略与决策树的结合方式。",
        "掌握 CART 决策树（分类与回归树）的构建过程，包括基尼系数（Gini Impurity）的计算与最优分裂点的选择。",
        "通过 Python 原生实现随机森林算法（不使用 sklearn 等第三方库的现成模型），加深对算法细节的理解。",
        "将随机森林算法应用于电商代运营场景下的 SKU 货品结构盘点分类问题，对比人工规则、单棵 CART 树与随机森林的效果。",
        "分析随机森林中特征重要性排序、树的数量对模型性能的影响，并通过可视化手段呈现实验结果。",
    ]
    for obj in objectives:
        add_bullet(doc, obj)

    # ================================================================
    # 二、实验原理
    # ================================================================
    add_heading_styled(doc, "二、实验原理", level=1)

    add_heading_styled(doc, "2.1 随机森林概述", level=2)
    add_body_text(doc,
        "随机森林（Random Forest）是一种基于 Bagging（Bootstrap Aggregating）策略的集成学习方法。"
        "它通过构建多棵决策树并对其预测结果进行投票（分类任务）或平均（回归任务）来提升模型的泛化能力与稳定性。"
        "随机森林的核心思想在于\"集成多样性\"——通过在数据样本和特征两个维度引入随机性，确保每棵树之间存在差异，"
        "从而降低过拟合风险并提高整体预测性能。"
    )

    add_heading_styled(doc, "2.2 CART 决策树", level=2)
    add_body_text(doc,
        "CART（Classification and Regression Tree）是一种二叉树结构的决策树算法。对于分类任务，CART 使用基尼系数"
        "（Gini Impurity）作为节点分裂的衡量标准。基尼系数越小，表示数据集的纯度越高。"
    )
    add_body_text(doc,
        "基尼系数的计算公式为：\n"
        "    Gini(D) = 1 - Σ(p_i)²\n"
        "其中 p_i 为数据集中第 i 类样本所占的比例。\n\n"
        "在节点分裂时，算法遍历所有特征及所有可能的切分点，选择使得分裂后加权基尼系数最小的特征和阈值作为最优分裂点。"
        "分裂后的基尼系数为：\n"
        "    Gini_split(D, A) = (|D₁|/|D|)×Gini(D₁) + (|D₂|/|D|)×Gini(D₂)"
    )

    add_heading_styled(doc, "2.3 Bagging 集成策略", level=2)
    add_body_text(doc,
        "Bagging（Bootstrap Aggregating）通过在原始训练集中进行有放回抽样（Bootstrap Sampling）生成多个不同的训练子集，"
        "然后分别训练基学习器。随机森林在 Bagging 的基础上进一步引入了随机子空间（Random Subspace）方法："
        "在每个节点分裂时，不是考虑所有特征，而是随机选择一个特征子集（通常为 √p 个，p 为总特征数），"
        "从中选择最优分裂特征。"
    )

    add_heading_styled(doc, "2.4 特征重要性评估", level=2)
    add_body_text(doc,
        "随机森林可以基于基尼系数减少量（Gini Importance）评估每个特征的重要性。"
        "对于每棵树，计算每个特征在分裂时带来的基尼系数减少量，并将减少量按特征累加后归一化，即得到各特征的重要性评分。"
        "特征重要性越高，说明该特征对分类决策的贡献越大。"
    )

    # ================================================================
    # 三、实验环境
    # ================================================================
    add_heading_styled(doc, "三、实验环境", level=1)
    env_data = [
        ("编程语言", "Python 3.10"),
        ("运行平台", "Ubuntu (WSL2)"),
        ("核心依赖库", "NumPy, Pandas, Matplotlib, Scikit-learn（仅用于评估指标）"),
        ("开发工具", "VS Code"),
    ]
    table = doc.add_table(rows=len(env_data), cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(env_data):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(11)

    doc.add_paragraph()

    # ================================================================
    # 四、数据集说明
    # ================================================================
    add_heading_styled(doc, "四、数据集说明", level=1)

    add_heading_styled(doc, "4.1 业务背景", level=2)
    add_body_text(doc,
        "本实验模拟电商代运营场景下的 SKU 货品结构盘点分类问题。电商代运营公司需要根据商品的表现数据，"
        "将 SKU 划分为不同的货品类别，以制定差异化的运营策略。合理地进行货品分类是电商精细化运营的基础。"
    )

    add_heading_styled(doc, "4.2 数据特征", level=2)
    add_body_text(doc, "数据集共包含 200 条 SKU 记录，9 个特征维度和 1 个目标标签，具体如下：")

    feature_table = doc.add_table(rows=10, cols=3)
    feature_table.style = "Light Grid Accent 1"
    feature_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["序号", "特征名称", "特征说明"]
    for j, h in enumerate(headers):
        cell = feature_table.cell(0, j)
        cell.text = h
        set_cell_shading(cell, HEADER_BG)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True

    features = [
        ("1", "近30天访客数", "商品近30天的访客数量"),
        ("2", "点击率", "商品详情页的点击转化率"),
        ("3", "30天销量", "近30天的销售数量"),
        ("4", "动销率", "商品动销比例"),
        ("5", "毛利率", "商品的毛利率"),
        ("6", "客单价", "平均每个客户带来的交易金额"),
        ("7", "库存周转天数", "库存周转一次所需的天数"),
        ("8", "现货库存", "当前现货库存数量"),
        ("9", "投放ROI", "广告投放的投入产出比"),
    ]
    for i, (idx, name, desc) in enumerate(features):
        feature_table.cell(i + 1, 0).text = idx
        feature_table.cell(i + 1, 1).text = name
        feature_table.cell(i + 1, 2).text = desc
        for cell in feature_table.rows[i + 1].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    add_heading_styled(doc, "4.3 标签类别", level=2)
    class_table = doc.add_table(rows=5, cols=3)
    class_table.style = "Light Grid Accent 1"
    class_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    class_headers = ["标签值", "类别名称", "业务含义"]
    for j, h in enumerate(class_headers):
        cell = class_table.cell(0, j)
        cell.text = h
        set_cell_shading(cell, HEADER_BG)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True

    classes = [
        ("0", "引流款", "高访客、高销量，用于吸引流量"),
        ("1", "利润款", "中等销量、高毛利，贡献利润"),
        ("2", "形象款", "低销量、高客单价，提升品牌形象"),
        ("3", "滞销清库存款", "低访客、低销量，需清仓处理"),
    ]
    for i, (idx, name, desc) in enumerate(classes):
        class_table.cell(i + 1, 0).text = idx
        class_table.cell(i + 1, 1).text = name
        class_table.cell(i + 1, 2).text = desc
        for cell in class_table.rows[i + 1].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    add_heading_styled(doc, "4.4 数据分布", level=2)
    add_body_text(doc,
        "数据集的标签分布为：引流款 58 个（29.0%）、利润款 61 个（30.5%）、"
        "形象款 27 个（13.5%）、滞销清库存款 54 个（27.0%）。"
        "训练集 140 条（70%），测试集 60 条（30%），采用分层抽样保持各类别比例一致。"
    )
    add_image_centered(
        doc, FIGURES_DIR / "rf_class_distribution.png",
        width_inches=4.5, caption="图1：SKU 类别分布"
    )

    # ================================================================
    # 五、算法实现
    # ================================================================
    add_heading_styled(doc, "五、算法实现", level=1)

    add_heading_styled(doc, "5.1 CART 决策树实现 (cart_base.py)", level=2)
    add_body_text(doc,
        "CART 决策树是本实验的基础组件，核心函数包括："
    )
    cart_funcs = [
        "gini(y)：计算数据集的基尼系数",
        "best_split(X, y)：遍历所有特征与切分点，寻找最优分裂",
        "build_cart_tree(X, y, max_depth, ...)：递归构建 CART 决策树",
        "predict(tree, X)：使用训练好的决策树进行预测",
        "format_rules(tree, ...)：将决策树以可读的 IF-THEN 规则形式输出",
    ]
    for f in cart_funcs:
        add_bullet(doc, f)

    add_body_text(doc, "关键实现细节如下：")
    add_code_block(doc,
        'def best_split(X, y, min_samples_leaf=2):\n'
        '    for fid in range(n_features):\n'
        '        thresholds = (unique[:-1] + unique[1:]) / 2.0\n'
        '        for t in thresholds:\n'
        '            X_left, y_left, X_right, y_right = split_dataset(X, y, fid, t)\n'
        '            gain = current_gini - (w_left * gini_left + w_right * gini_right)\n'
        '            if gain > best_gain: best_feature, best_threshold, best_gain = fid, t, gain'
    )

    add_heading_styled(doc, "5.2 随机森林实现 (random_forest.py)", level=2)
    add_body_text(doc,
        "RandomForest 类的核心实现包括："
    )
    rf_funcs = [
        "bootstrap_sample(X, y, rng)：对训练数据进行有放回 Bootstrap 抽样",
        "fit(X, y)：训练随机森林，每棵树在随机特征子集上使用 CART 算法训练",
        "predict(X)：集成所有树的预测结果，采用多数投票法",
        "predict_proba(X)：输出每个样本属于各类别的概率估计",
        "_gini_importance(tree)：递归计算树中各特征的基尼重要性",
    ]
    for f in rf_funcs:
        add_bullet(doc, f)

    add_body_text(doc, "超参数配置：")
    params = [
        ("n_estimators", "15", "决策树的数量"),
        ("max_depth", "4", "每棵树的最大深度"),
        ("min_samples_split", "5", "内部节点再划分所需的最小样本数"),
        ("min_samples_leaf", "2", "叶节点的最少样本数"),
        ("max_features", "sqrt（默认）", "每棵树随机选择的特征数上限"),
    ]
    param_table = doc.add_table(rows=len(params) + 1, cols=3)
    param_table.style = "Light Grid Accent 1"
    param_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    param_headers = ["参数名", "取值", "说明"]
    for j, h in enumerate(param_headers):
        cell = param_table.cell(0, j)
        cell.text = h
        set_cell_shading(cell, HEADER_BG)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
    for i, (k, v, d) in enumerate(params):
        param_table.cell(i + 1, 0).text = k
        param_table.cell(i + 1, 1).text = v
        param_table.cell(i + 1, 2).text = d
        for cell in param_table.rows[i + 1].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # ================================================================
    # 六、实验内容与步骤
    # ================================================================
    add_heading_styled(doc, "六、实验内容与步骤", level=1)

    add_heading_styled(doc, "6.1 实验流程", level=2)
    steps = [
        "数据加载与预处理：读取 CSV 数据，填充缺失值，划分训练集与测试集（70%/30%）。",
        "人工规则分类：基于业务经验设定简单的分类规则作为基准。",
        "单棵 CART 决策树：使用实现好的 CART 算法训练一棵决策树。",
        "随机森林训练：使用 Bootstrap 抽样 + 随机特征子空间训练 15 棵 CART 树。",
        "模型评估：在测试集上计算准确率、精确率、召回率、F1-score；输出混淆矩阵。",
        "特征重要性分析：基于基尼系数减少量排名各特征的重要性。",
        "超参数影响分析：测试不同 n_estimators 对模型性能的影响。",
        "全店货品结构盘点：使用训练好的随机森林对全量数据进行分类，输出货品结构分布。",
    ]
    for i, step in enumerate(steps):
        add_bullet(doc, step)

    add_heading_styled(doc, "6.2 人工规则设计", level=2)
    add_body_text(doc,
        "基于电商运营的业务经验设计人工分类规则作为对比基准："
    )
    add_code_block(doc,
        "若 30天销量 > 200 → 引流款\n"
        "否则若 毛利率 > 0.30 → 利润款\n"
        "否则若 毛利率 > 0.50 且 30天销量 < 30 → 形象款\n"
        "否则 → 滞销清库存款"
    )

    add_heading_styled(doc, "6.3 CART 决策树规则", level=2)
    add_body_text(doc,
        "实验训练出的 CART 决策树（max_depth=4）提取的分类规则如下："
    )
    add_code_block(doc,
        "若 近30天访客数 <= 736.0\n"
        "    若 近30天访客数 <= 200.5\n"
        "        若 毛利率 <= 0.4219 → 滞销清库存款\n"
        "        若 毛利率 > 0.4219  → 形象款\n"
        "    若 近30天访客数 > 200.5 → 利润款\n"
        "若 近30天访客数 > 736.0 → 引流款"
    )
    add_body_text(doc,
        "该决策树仅使用了\"近30天访客数\"和\"毛利率\"两个特征，深度为 3（根节点深度为 1），"
        "共有 4 个叶节点（决策规则），结构非常简洁且可解释性强。"
    )

    # ================================================================
    # 七、实验结果与分析
    # ================================================================
    add_heading_styled(doc, "七、实验结果与分析", level=1)

    add_heading_styled(doc, "7.1 三种方法对比", level=2)
    add_body_text(doc, "下表汇总了三种分类方法的性能对比：")

    comp_table = doc.add_table(rows=4, cols=4)
    comp_table.style = "Light Grid Accent 1"
    comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    comp_headers = ["方法", "训练准确率", "测试准确率", "耗时(s)"]
    for j, h in enumerate(comp_headers):
        cell = comp_table.cell(0, j)
        cell.text = h
        set_cell_shading(cell, HEADER_BG)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
    comp_data = [
        ("人工规则", "0.8500", "0.8333", "0.0001"),
        ("单棵 CART 树", "1.0000", "1.0000", "0.0457"),
        ("随机森林 (n=15)", "1.0000", "1.0000", "0.1550"),
    ]
    for i, row in enumerate(comp_data):
        for j, v in enumerate(row):
            comp_table.cell(i + 1, j).text = v
            for paragraph in comp_table.cell(i + 1, j).paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    add_body_text(doc,
        "分析：人工规则的测试准确率为 83.33%，作为基线方法表现尚可，但无法识别形象款（精确率与召回率均为 0），"
        "因为简单规则难以捕捉复杂的分类边界。单棵 CART 树和随机森林在测试集上均达到了 100% 的准确率，"
        "说明数据在 9 个特征维度上分类边界清晰，CART 树已能完美拟合。随机森林由于需要训练 15 棵树，"
        "训练耗时（0.1550s）约为单棵 CART 树（0.0457s）的 3.4 倍，但预测速度仍然很快（0.0016s）。"
    )

    add_heading_styled(doc, "7.2 混淆矩阵分析", level=2)

    add_image_centered(
        doc, FIGURES_DIR / "rf_confusion_manual.png",
        width_inches=4.0, caption="图2：人工规则混淆矩阵"
    )
    add_image_centered(
        doc, FIGURES_DIR / "rf_confusion_cart.png",
        width_inches=4.0, caption="图3：CART 决策树混淆矩阵"
    )
    add_image_centered(
        doc, FIGURES_DIR / "rf_confusion_rf.png",
        width_inches=4.0, caption="图4：随机森林混淆矩阵"
    )

    add_body_text(doc,
        "从混淆矩阵可以直观看出：人工规则将 8 个实际为形象款的样本全部误判（主要被分到利润款和滞销清库存款），"
        "而 CART 树和随机森林均正确分类了所有样本。"
    )

    add_heading_styled(doc, "7.3 分类报告对比", level=2)

    add_body_text(doc, "人工规则分类报告：", bold=True)
    add_code_block(doc,
        "              精度    召回率    F1分数    样本数\n"
        "引流款        1.00    0.89     0.94      18\n"
        "利润款        0.69    1.00     0.82      18\n"
        "形象款        0.00    0.00     0.00       8\n"
        "滞销清库存款    0.89    1.00     0.94      16\n"
        "宏观平均       0.65    0.72     0.68      60\n"
        "加权平均       0.74    0.83     0.78      60"
    )

    add_body_text(doc, "CART 决策树分类报告：", bold=True)
    add_code_block(doc,
        "              精度    召回率    F1分数    样本数\n"
        "引流款        1.00    1.00     1.00      18\n"
        "利润款        1.00    1.00     1.00      18\n"
        "形象款        1.00    1.00     1.00       8\n"
        "滞销清库存款    1.00    1.00     1.00      16\n"
        "宏观平均       1.00    1.00     1.00      60\n"
        "加权平均       1.00    1.00     1.00      60"
    )

    add_body_text(doc, "随机森林分类报告：", bold=True)
    add_code_block(doc,
        "              精度    召回率    F1分数    样本数\n"
        "引流款        1.00    1.00     1.00      18\n"
        "利润款        1.00    1.00     1.00      18\n"
        "形象款        1.00    1.00     1.00       8\n"
        "滞销清库存款    1.00    1.00     1.00      16\n"
        "宏观平均       1.00    1.00     1.00      60\n"
        "加权平均       1.00    1.00     1.00      60"
    )

    add_body_text(doc,
        "CART 与随机森林在所有类别上均取得 1.00 的精度/召回率/F1，说明该数据集线性可分。"
        "人工规则的形象款精确率和召回率均为 0，说明经验规则对该类别完全失效。"
    )

    add_heading_styled(doc, "7.4 随机森林评估指标表", level=2)
    add_image_centered(
        doc, FIGURES_DIR / "rf_metrics_table.png",
        width_inches=5.5, caption="图5：随机森林详细评估指标"
    )

    add_heading_styled(doc, "7.5 特征重要性分析", level=2)
    add_image_centered(
        doc, FIGURES_DIR / "rf_feature_importance.png",
        width_inches=5.0, caption="图6：随机森林特征重要性排序"
    )

    importance_data = [
        ("1", "30天销量", "0.2012"),
        ("2", "近30天访客数", "0.1736"),
        ("3", "客单价", "0.1610"),
        ("4", "毛利率", "0.1143"),
        ("5", "投放ROI", "0.1032"),
        ("6", "现货库存", "0.0972"),
        ("7", "库存周转天数", "0.0602"),
        ("8", "动销率", "0.0477"),
        ("9", "点击率", "0.0416"),
    ]
    imp_table = doc.add_table(rows=len(importance_data) + 1, cols=3)
    imp_table.style = "Light Grid Accent 1"
    imp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    imp_headers = ["排名", "特征", "重要性分数"]
    for j, h in enumerate(imp_headers):
        cell = imp_table.cell(0, j)
        cell.text = h
        set_cell_shading(cell, HEADER_BG)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
    for i, (rank, name, score) in enumerate(importance_data):
        imp_table.cell(i + 1, 0).text = rank
        imp_table.cell(i + 1, 1).text = name
        imp_table.cell(i + 1, 2).text = score
        for cell in imp_table.rows[i + 1].cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    add_body_text(doc,
        "分析：\"30天销量\"（0.2012）、\"近30天访客数\"（0.1736）和\"客单价\"（0.1610）是最重要的三个特征，"
        "三者累计重要性超过 50%。这与电商业务常识一致——销量和访客数是区分引流款与其他款式的关键指标，"
        "客单价则是区分利润款/形象款的重要依据。\"点击率\"（0.0416）和\"动销率\"（0.0477）重要性最低，"
        "说明这两个特征在当前分类任务中的区分能力有限。"
    )

    add_heading_styled(doc, "7.6 树数量对模型性能的影响", level=2)
    add_image_centered(
        doc, FIGURES_DIR / "rf_estimator_comparison.png",
        width_inches=5.0, caption="图7：不同 n_estimators 下随机森林性能对比"
    )
    add_body_text(doc,
        "从图中可以看出：即使只有 1 棵树，由于随机森林使用了随机特征子空间，其分类性能仍然很好。"
        "当 n_estimators 增加到 3 棵时，模型已趋于稳定。"
        "随着树的数量继续增加，训练集和测试集准确率均保持在 1.0 的水平，没有出现明显的过拟合。"
        "这验证了随机森林的集成策略在数据质量较好时能够稳定地提升或保持模型性能。"
    )

    add_heading_styled(doc, "7.7 全店货品结构分布", level=2)
    add_body_text(doc,
        "使用训练好的随机森林对全量 200 条 SKU 数据进行分类，得到全店货品结构分布如下："
    )
    structure_data = [
        ("引流款", "58", "29.0%"),
        ("利润款", "61", "30.5%"),
        ("形象款", "27", "13.5%"),
        ("滞销清库存款", "54", "27.0%"),
        ("合计", "200", "100%"),
    ]
    struct_table = doc.add_table(rows=len(structure_data) + 1, cols=3)
    struct_table.style = "Light Grid Accent 1"
    struct_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    s_headers = ["类别", "SKU 数量", "占比"]
    for j, h in enumerate(s_headers):
        cell = struct_table.cell(0, j)
        cell.text = h
        set_cell_shading(cell, HEADER_BG)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
    for i, (cls, cnt, pct) in enumerate(structure_data):
        struct_table.cell(i + 1, 0).text = cls
        struct_table.cell(i + 1, 1).text = cnt
        struct_table.cell(i + 1, 2).text = pct
        for cell in struct_table.rows[i + 1].cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    add_body_text(doc,
        "从货品结构来看，引流款和利润款合计占比 59.5%，是店铺的主力商品；滞销清库存款占 27.0%，"
        "比例偏高，建议店铺运营重点关注库存清理和商品汰换策略。形象款占 13.5%，比例适中，"
        "有助于提升店铺整体品牌形象。"
    )

    # ================================================================
    # 八、结论与体会
    # ================================================================
    add_heading_styled(doc, "八、结论与体会", level=1)

    conclusions = [
        "通过本次实验，成功使用 Python 原生实现了随机森林算法，包括 CART 决策树的构建、"
        "Bootstrap 抽样、随机特征子空间、多数投票等核心机制，加深了对集成学习原理的理解。",

        "在电商代运营的 SKU 货品结构盘点分类场景中，随机森林和单棵 CART 树均取得了 100% 的测试准确率，"
        "说明该数据集在 9 个特征构成的特征空间中具有清晰的分类边界。人工规则（83.3%）虽然简单实用，"
        "但在识别形象款等边界模糊的类别时存在明显不足。",

        "特征重要性分析表明，\"30天销量\"、\"近30天访客数\"和\"客单价\"是最关键的三个分类特征，"
        "这与电商运营的业务直觉高度一致——高销量高访客的商品为引流款，高毛利适中的商品为利润款，"
        "高客单价低销量的商品为形象款，各项指标均低的商品为滞销清库存款。",

        "CART 决策树提取的可解释规则（仅用两个特征即可完成分类）具有很高的业务落地价值，"
        "运营人员可以直接根据\"访客数\"和\"毛利率\"两个指标快速对商品进行分类。",

        "随机森林相比单棵树虽然在准确率上没有进一步提升（数据已完美可分），但其集成特性"
        "使其对噪声和异常值更鲁棒，且能提供特征重要性和概率估计等额外信息。",
    ]
    for i, c in enumerate(conclusions):
        add_body_text(doc, f"{i+1}. {c}")

    # ========== 保存 ==========
    doc.save(str(REPORT_PATH))
    print(f"[完成] 实验报告已保存至：{REPORT_PATH}")


if __name__ == "__main__":
    build_report()
